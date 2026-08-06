import os
from typing import Dict, Any
from app.ingestion.base_processor import BaseFileProcessor

class DocxProcessor(BaseFileProcessor):
    def can_handle(self, file_extension: str) -> bool:
        return file_extension.lower() in ['.docx', '.doc']

    def extract_content(self, file_path: str, offline_mode: bool = False) -> Dict[str, Any]:
        filename = os.path.basename(file_path)
        try:
            import docx
            doc = docx.Document(file_path)
            lines = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Extract text inside Word document tables
            table_lines = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        table_lines.append(row_text)
                        
            if table_lines:
                lines.append("\n[Document Tables]:")
                lines.extend(table_lines)
                
            full_text = "\n".join(lines)
            return {
                "content": full_text,
                "engine_used": "python_docx_extractor",
                "is_offline": True,
                "warning": None
            }
        except ImportError:
            # Fallback if python-docx is missing
            return {
                "content": f"[Word Document {filename} attached (python-docx not installed)]",
                "engine_used": "raw_docx_stub",
                "is_offline": True,
                "warning": "python-docx library not installed"
            }
        except Exception as e:
            return {
                "content": f"[!] Error extracting Word document ({filename}): {str(e)}",
                "engine_used": "python_docx_failed",
                "is_offline": True,
                "warning": str(e)
            }

    @property
    def output_key(self) -> str:
        return "text"
